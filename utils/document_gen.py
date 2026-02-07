"""
Módulo para generación de documentos Word (actas)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import streamlit as st


def generate_word_document(analysis, meeting_info, transcription=""):
    """
    Genera un documento Word con el acta de la reunión en formato institucional
    
    Args:
        analysis: Diccionario con el análisis estructurado
        meeting_info: Información de la reunión
        transcription: Transcripción completa (opcional)
        
    Returns:
        Document: Objeto documento de python-docx
    """
    try:
        doc = Document()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # ENCABEZADO INSTITUCIONAL
        add_institutional_header(doc, meeting_info)
        
        # ASISTENTES
        if meeting_info.get("asistentes"):
            add_asistentes_section(doc, meeting_info["asistentes"])
        
        # AGENDA
        if meeting_info.get("agenda"):
            add_agenda_section(doc, meeting_info["agenda"])
        
        # DESARROLLO DE LA REUNIÓN
        if analysis.get("desarrollo"):
            add_desarrollo_section(doc, analysis["desarrollo"])
        
        # DECISIONES TOMADAS
        if analysis.get("decisiones"):
            add_list_section(doc, "DECISIONES TOMADAS", analysis["decisiones"])
        
        # TAREAS Y RESPONSABLES
        if analysis.get("tareas"):
            add_table_section(doc, "TAREAS Y RESPONSABLES", analysis["tareas"])
        
        # PRÓXIMOS PASOS
        if analysis.get("proximos_pasos"):
            add_list_section(doc, "PRÓXIMOS PASOS", analysis["proximos_pasos"])
        
        # ANEXO: Transcripción completa (opcional)
        if transcription:
            add_transcription_section(doc, transcription)
        
        return doc
        
    except Exception as e:
        st.error(f"Error al generar documento: {str(e)}")
        return None


def add_institutional_header(doc, meeting_info):
    """Agrega el encabezado institucional del acta"""
    # Título principal
    title = doc.add_heading(f"ACTA No. {meeting_info.get('numero_acta', '___')}", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # Nombre del comité
    comite = doc.add_paragraph(meeting_info.get("comite", ""))
    comite.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    comite_run = comite.runs[0]
    comite_run.font.size = Pt(12)
    comite_run.font.bold = True
    
    doc.add_paragraph()  # Espacio
    
    # Información en tabla
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Datos
    data = [
        ("Área que convoca y organiza:", meeting_info.get("area_convoca", "")),
        ("Fecha de realización:", meeting_info.get("fecha", "")),
        ("Hora de inicio:", meeting_info.get("hora_inicio", "")),
        ("Hora de finalización:", meeting_info.get("hora_fin", "")),
        ("Lugar:", meeting_info.get("lugar", "")),
        ("Notas tomadas por:", meeting_info.get("notas_por", "")),
        ("", "")  # Fila vacía
    ]
    
    for idx, (label, value) in enumerate(data):
        row = table.rows[idx]
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        
        cell_label.text = label
        cell_value.text = value
        
        # Formato
        for cell in [cell_label, cell_value]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    if cell == cell_label and label:
                        run.font.bold = True
    
    doc.add_paragraph()  # Espacio


def add_asistentes_section(doc, asistentes):
    """Agrega la sección de asistentes"""
    heading = doc.add_heading("ASISTENTES", level=1)
    heading_format = heading.runs[0].font
    heading_format.color.rgb = RGBColor(31, 78, 121)
    heading_format.size = Pt(14)
    
    # Tabla de asistentes
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Encabezados
    header_cells = table.rows[0].cells
    header_cells[0].text = "Nombre"
    header_cells[1].text = "Cargo"
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
                run.font.size = Pt(11)
    
    # Agregar asistentes
    for asistente in asistentes:
        row_cells = table.add_row().cells
        row_cells[0].text = asistente.get("nombre", "")
        row_cells[1].text = asistente.get("cargo", "")
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    
    doc.add_paragraph()  # Espacio


def add_agenda_section(doc, agenda_text):
    """Agrega la sección de agenda"""
    heading = doc.add_heading("AGENDA", level=1)
    heading_format = heading.runs[0].font
    heading_format.color.rgb = RGBColor(31, 78, 121)
    heading_format.size = Pt(14)
    
    # Dividir agenda en líneas
    agenda_lines = agenda_text.strip().split('\n')
    
    for line in agenda_lines:
        line = line.strip()
        if line:
            para = doc.add_paragraph(line, style='List Number')
            para.paragraph_format.left_indent = Inches(0.25)
            for run in para.runs:
                run.font.size = Pt(11)
    
    doc.add_paragraph()  # Espacio


def add_desarrollo_section(doc, desarrollo_text):
    """Agrega la sección de desarrollo de la reunión"""
    heading = doc.add_heading("DESARROLLO DE LA REUNIÓN", level=1)
    heading_format = heading.runs[0].font
    heading_format.color.rgb = RGBColor(31, 78, 121)
    heading_format.size = Pt(14)
    
    para = doc.add_paragraph(desarrollo_text)
    para_format = para.paragraph_format
    para_format.space_after = Pt(12)
    para_format.line_spacing = 1.15
    para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    for run in para.runs:
        run.font.size = Pt(11)
    
    doc.add_paragraph()  # Espacio


def add_list_section(doc, title, items):
    """Agrega una sección con lista de items"""
    # Título de sección
    heading = doc.add_heading(title, level=1)
    heading_format = heading.runs[0].font
    heading_format.color.rgb = RGBColor(31, 78, 121)
    
    # Items
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.25)


def add_table_section(doc, title, tasks):
    """Agrega una sección con tabla de tareas"""
    # Título de sección
    heading = doc.add_heading(title, level=1)
    heading_format = heading.runs[0].font
    heading_format.color.rgb = RGBColor(31, 78, 121)
    
    # Crear tabla
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Encabezados
    header_cells = table.rows[0].cells
    headers = ['Tarea', 'Responsable', 'Fecha Límite']
    
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        # Negrita para encabezados
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
    
    # Agregar tareas
    for task in tasks:
        row_cells = table.add_row().cells
        
        # Intentar parsear tarea en formato "Tarea | Responsable | Fecha"
        parts = task.split('|')
        if len(parts) >= 3:
            row_cells[0].text = parts[0].strip()
            row_cells[1].text = parts[1].strip()
            row_cells[2].text = parts[2].strip()
        elif len(parts) == 2:
            row_cells[0].text = parts[0].strip()
            row_cells[1].text = parts[1].strip()
            row_cells[2].text = "Por definir"
        else:
            row_cells[0].text = task
            row_cells[1].text = "Por asignar"
            row_cells[2].text = "Por definir"
    
    doc.add_paragraph()  # Espacio


def add_transcription_section(doc, transcription):
    """Agrega la transcripción completa como anexo"""
    doc.add_page_break()
    
    # Título
    heading = doc.add_heading("ANEXO: TRANSCRIPCIÓN COMPLETA", level=1)
    heading_format = heading.runs[0].font
    heading_format.color.rgb = RGBColor(31, 78, 121)
    
    # Transcripción
    para = doc.add_paragraph(transcription)
    para_format = para.paragraph_format
    para_format.line_spacing = 1.0
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)


def save_document(doc, filename="acta_reunion.docx"):
    """
    Guarda el documento en disco
    
    Args:
        doc: Documento de python-docx
        filename: Nombre del archivo
        
    Returns:
        str: Ruta del archivo guardado
    """
    try:
        doc.save(filename)
        return filename
    except Exception as e:
        st.error(f"Error al guardar documento: {str(e)}")
        return None
